import os
import sys
import argparse
from pathlib import Path
from typing import Optional, List
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DocumentConverter:
    """Convert various document formats to Markdown (excluding PDF)."""
    
    def __init__(self):
        """Initialize converter and check dependencies."""
        self.supported_formats = {
            '.docx': self._convert_docx,
            '.doc': self._convert_doc,
            '.html': self._convert_html,
            '.htm': self._convert_html,
            '.txt': self._convert_txt,
            '.rtf': self._convert_rtf,
        }
    
    def convert(self, input_path: str, output_path: Optional[str] = None) -> bool:
        """
        Convert a document to Markdown.
        
        Args:
            input_path: Path to input document
            output_path: Path to output Markdown file (optional)
        
        Returns:
            True if conversion successful, False otherwise
        """
        input_file = Path(input_path)
        
        if not input_file.exists():
            logger.error(f"Input file not found: {input_path}")
            return False
        
        # Check if format is supported
        ext = input_file.suffix.lower()
        
        # Redirect PDF files to dedicated converter
        if ext == '.pdf':
            logger.warning("For PDF conversion, please use: python pdf_to_md_converter.py")
            logger.warning(f"Command: python pdf_to_md_converter.py {input_path}")
            return False
        
        if ext not in self.supported_formats:
            logger.error(f"Unsupported format: {ext}")
            logger.info(f"Supported formats: {', '.join(self.supported_formats.keys())}")
            logger.info("For PDF files, use pdf_to_md_converter.py")
            return False
        
        # Determine output path
        if output_path is None:
            output_path = input_file.with_suffix('.md')
        else:
            output_path = Path(output_path)
        
        logger.info(f"Converting {input_file.name} to Markdown...")
        
        try:
            # Call appropriate conversion method
            converter_func = self.supported_formats[ext]
            markdown_content = converter_func(input_file)
            
            # Add metadata header
            metadata = self._create_metadata(input_file)
            full_content = metadata + markdown_content
            
            # Write output
            output_path.write_text(full_content, encoding='utf-8')
            logger.info(f"Successfully converted to: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Conversion failed: {str(e)}")
            import traceback
            logger.debug(traceback.format_exc())
            return False
    
    def _create_metadata(self, input_file: Path) -> str:
        """Create metadata header for the Markdown file."""
        return f"""# {input_file.stem}

*Converted from: {input_file.name}*

---

"""
    
    def _convert_docx(self, input_file: Path) -> str:
        """Convert DOCX to Markdown using python-docx."""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
        
        doc = Document(input_file)
        markdown_lines = []
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_lines.append('')
                continue
            
            # Handle different paragraph styles
            style = paragraph.style.name.lower()
            
            if 'heading 1' in style:
                markdown_lines.append(f"# {text}")
            elif 'heading 2' in style:
                markdown_lines.append(f"## {text}")
            elif 'heading 3' in style:
                markdown_lines.append(f"### {text}")
            elif 'heading 4' in style:
                markdown_lines.append(f"#### {text}")
            elif 'heading 5' in style:
                markdown_lines.append(f"##### {text}")
            elif 'heading 6' in style:
                markdown_lines.append(f"###### {text}")
            elif 'list' in style:
                # Handle lists
                if paragraph.style.name.startswith('List Number'):
                    markdown_lines.append(f"1. {text}")
                else:
                    markdown_lines.append(f"- {text}")
            else:
                # Handle bold and italic text
                formatted_text = self._format_runs(paragraph)
                markdown_lines.append(formatted_text)
            
            markdown_lines.append('')
        
        # Process tables
        for table in doc.tables:
            markdown_lines.append('')
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                markdown_lines.append('| ' + ' | '.join(cells) + ' |')
                
                # Add separator after header row
                if i == 0:
                    markdown_lines.append('|' + '|'.join(['---'] * len(cells)) + '|')
            markdown_lines.append('')
        
        return '\n'.join(markdown_lines)
    
    def _format_runs(self, paragraph) -> str:
        """Format text runs with bold and italic."""
        result = []
        for run in paragraph.runs:
            text = run.text
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            result.append(text)
        return ''.join(result)
    
    def _convert_doc(self, input_file: Path) -> str:
        """Convert DOC to Markdown using pypandoc."""
        try:
            import pypandoc
        except ImportError:
            logger.error("pypandoc not installed. Run: pip install pypandoc")
            raise
        
        return pypandoc.convert_file(str(input_file), 'md')
    
    def _convert_html(self, input_file: Path) -> str:
        """Convert HTML to Markdown using html2text."""
        try:
            import html2text
        except ImportError:
            logger.error("html2text not installed. Run: pip install html2text")
            raise
        
        html_content = input_file.read_text(encoding='utf-8')
        
        # Configure html2text
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.ignore_emphasis = False
        h.body_width = 0  # Don't wrap lines
        h.unicode_snob = True
        h.skip_internal_links = False
        
        return h.handle(html_content)
    
    def _convert_txt(self, input_file: Path) -> str:
        """Convert plain text to Markdown (with smart formatting)."""
        content = input_file.read_text(encoding='utf-8')
        
        lines = content.split('\n')
        markdown_lines = []
        in_list = False
        
        for line in lines:
            stripped = line.strip()
            
            # Empty line
            if not stripped:
                markdown_lines.append('')
                in_list = False
                continue
            
            # Detect potential headers (all caps, short lines)
            if stripped.isupper() and len(stripped) < 100 and not any(c.isdigit() for c in stripped):
                markdown_lines.append(f"## {stripped.title()}")
                in_list = False
            # Detect numbered lists
            elif stripped.startswith(tuple(f"{i}. " for i in range(1, 100))):
                markdown_lines.append(line)
                in_list = True
            # Detect bullet points
            elif stripped.startswith(('- ', '* ', '+ ', '• ')):
                if stripped.startswith('• '):
                    markdown_lines.append('- ' + stripped[2:])
                else:
                    markdown_lines.append(line)
                in_list = True
            else:
                # Regular text
                if in_list:
                    markdown_lines.append('  ' + line)  # Indent continuation
                else:
                    markdown_lines.append(line)
        
        return '\n'.join(markdown_lines)
    
    def _convert_rtf(self, input_file: Path) -> str:
        """Convert RTF to Markdown using pypandoc."""
        try:
            import pypandoc
        except ImportError:
            logger.error("pypandoc not installed. Run: pip install pypandoc")
            raise
        
        return pypandoc.convert_file(str(input_file), 'md')
    
    def batch_convert(self, input_dir: str, output_dir: Optional[str] = None) -> List[str]:
        """
        Convert all supported documents in a directory.
        
        Args:
            input_dir: Directory containing documents
            output_dir: Output directory for Markdown files
        
        Returns:
            List of successfully converted files
        """
        input_path = Path(input_dir)
        
        if not input_path.is_dir():
            logger.error(f"Not a directory: {input_dir}")
            return []
        
        if output_dir:
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
        else:
            output_path = input_path
        
        converted_files = []
        
        # Find all supported files
        for ext in self.supported_formats.keys():
            for file in input_path.glob(f"*{ext}"):
                output_file = output_path / file.with_suffix('.md').name
                
                if self.convert(str(file), str(output_file)):
                    converted_files.append(str(file))
        
        # Check for PDF files and remind user
        pdf_files = list(input_path.glob("*.pdf"))
        if pdf_files:
            logger.info(f"\nFound {len(pdf_files)} PDF file(s). Use pdf_to_md_converter.py to convert them:")
            for pdf_file in pdf_files[:5]:  # Show first 5
                logger.info(f"  python pdf_to_md_converter.py {pdf_file}")
            if len(pdf_files) > 5:
                logger.info(f"  ... and {len(pdf_files) - 5} more")
        
        logger.info(f"\nConverted {len(converted_files)} file(s)")
        return converted_files


def main():
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(
        description='Convert documents to Markdown format (DOCX, HTML, TXT, RTF)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert single file
  python doc_to_markdown.py document.docx
  
  # Convert with custom output
  python doc_to_markdown.py document.html output.md
  
  # Batch convert directory
  python doc_to_markdown.py ./documents --batch
  
  # Batch convert with output directory
  python doc_to_markdown.py ./documents --batch --output ./markdown

Note:
  For PDF files, use the dedicated converter:
    python pdf_to_md_converter.py document.pdf
        """
    )
    
    parser.add_argument(
        'input',
        help='Input file or directory'
    )
    
    parser.add_argument(
        'output',
        nargs='?',
        help='Output file or directory (optional)'
    )
    
    parser.add_argument(
        '--batch',
        action='store_true',
        help='Batch convert all files in directory'
    )
    
    parser.add_argument(
        '--verbose',
        action='store_true',
        help='Enable verbose logging'
    )
    
    args = parser.parse_args()
    
    if args.verbose:
        logger.setLevel(logging.DEBUG)
    
    converter = DocumentConverter()
    
    if args.batch:
        converter.batch_convert(args.input, args.output)
    else:
        success = converter.convert(args.input, args.output)
        sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()